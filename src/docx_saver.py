import re
from docx import Document
from docx.shared import Pt, RGBColor

# Configuration Constants
DEFAULT_FONT_NAME = 'Calibri'
DEFAULT_FONT_SIZE = 11

def save_markdown_to_docx(markdown_text: str, output_path: str):
    """
    Parses simple Markdown (Headers, Bold, Bullet points) and saves to Docx.
    """
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = Pt(DEFAULT_FONT_SIZE)

    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            apply_formatting(p, line[2:])
            
        # Normal text
        else:
            p = doc.add_paragraph()
            apply_formatting(p, line)
            
    doc.save(output_path)
    return output_path

def apply_formatting(paragraph, text):
    """
    Applies simple inline formatting (Bold only for now) to a paragraph object.
    Re-writes the paragraph content with runs.
    """
    paragraph.clear()
    
    # Simple bold parser: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)
